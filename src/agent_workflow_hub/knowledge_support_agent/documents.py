"""Deterministic document filtering, parsing, and chunking."""

from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
import hashlib
from pathlib import PurePosixPath
import re
from typing import Mapping


class DocumentError(ValueError):
    """Raised when source content cannot be safely parsed."""


class DocumentDependencyUnavailable(DocumentError):
    """Raised when an explicitly supported parser is not installed."""


class UnsupportedDocument(DocumentError):
    """Raised when a document needs an unsupported parser such as OCR."""


@dataclass(frozen=True, kw_only=True)
class KnowledgeChunk:
    chunk_id: str
    source_id: str
    source_kind: str
    source_version: str
    source_path: str
    title: str
    section: str
    content: str
    content_sha256: str
    provenance: Mapping[str, object]
    inferred: bool


_GIT_DOCUMENT_EXTENSIONS = {".md", ".markdown", ".txt"}
_CODE_EXTENSIONS = {".py", ".js", ".jsx", ".ts", ".tsx", ".vue"}
_EXCLUDED_EXTENSIONS = {
    ".env",
    ".har",
    ".ini",
    ".jpeg",
    ".jpg",
    ".log",
    ".png",
    ".trace",
    ".zip",
}
_EXCLUDED_PARTS = {
    ".git",
    ".idea",
    ".pytest_cache",
    ".venv",
    "__pycache__",
    "artifacts",
    "cache",
    "logs",
    "node_modules",
    "reports",
    "screenshots",
    "trace",
}
_EXCLUDED_NAMES = {
    ".env",
    "environment-validated.ini",
}
_HEADING = re.compile(r"^(#{1,6})\s+(.+?)\s*$")


def should_index_git_path(path: str, *, include_code: bool) -> bool:
    normalized = path.replace("\\", "/")
    pure = PurePosixPath(normalized)
    parts = tuple(part.casefold() for part in pure.parts)
    name = pure.name.casefold()
    suffix = pure.suffix.casefold()
    if (
        not normalized
        or normalized.startswith("/")
        or ".." in pure.parts
        or any(part in _EXCLUDED_PARTS for part in parts)
        or name in _EXCLUDED_NAMES
        or suffix in _EXCLUDED_EXTENSIONS
    ):
        return False
    if suffix in _GIT_DOCUMENT_EXTENSIONS:
        return True
    return include_code and suffix in _CODE_EXTENSIONS


def parse_document(
    *,
    source_id: str,
    source_kind: str,
    source_version: str,
    source_path: str,
    content: bytes,
    origin_url: str | None = None,
) -> tuple[KnowledgeChunk, ...]:
    suffix = PurePosixPath(source_path.replace("\\", "/")).suffix.casefold()
    if suffix in {".md", ".markdown"}:
        sections = _markdown_sections(_decode_text(content), source_path)
    elif suffix == ".txt" or not suffix:
        sections = [(PurePosixPath(source_path).name, _decode_text(content))]
    elif suffix == ".docx":
        sections = [(PurePosixPath(source_path).stem, _read_docx(content))]
    elif suffix == ".pdf":
        sections = [(PurePosixPath(source_path).stem, _read_pdf(content))]
    else:
        raise UnsupportedDocument(f"unsupported document type: {suffix or '<none>'}")

    title = PurePosixPath(source_path.replace("\\", "/")).name
    provenance: dict[str, object] = {"source_path": source_path}
    if source_kind == "git":
        provenance["commit_sha"] = source_version
    if origin_url is not None:
        provenance["origin_url"] = origin_url
    chunks: list[KnowledgeChunk] = []
    for section, text in sections:
        for index, part in enumerate(_split_text(text.strip())):
            chunks.append(
                _chunk(
                    source_id=source_id,
                    source_kind=source_kind,
                    source_version=source_version,
                    source_path=source_path,
                    title=title,
                    section=section,
                    content=part,
                    provenance={**provenance, "part": index},
                )
            )
    return tuple(chunks)


def _markdown_sections(text: str, source_path: str) -> list[tuple[str, str]]:
    headings: list[str] = []
    current_section = PurePosixPath(source_path).stem
    body: list[str] = []
    sections: list[tuple[str, str]] = []
    for line in text.splitlines():
        match = _HEADING.match(line)
        if match is None:
            body.append(line)
            continue
        if any(item.strip() for item in body):
            sections.append((current_section, "\n".join(body).strip()))
        level = len(match.group(1))
        headings = headings[: level - 1]
        headings.append(match.group(2).strip())
        current_section = " / ".join(headings)
        body = []
    if any(item.strip() for item in body):
        sections.append((current_section, "\n".join(body).strip()))
    return sections


def _split_text(text: str, maximum: int = 1200, overlap: int = 150) -> tuple[str, ...]:
    if not text:
        return ()
    if len(text) <= maximum:
        return (text,)
    step = maximum - overlap
    return tuple(text[start : start + maximum] for start in range(0, len(text), step))


def _chunk(
    *,
    source_id: str,
    source_kind: str,
    source_version: str,
    source_path: str,
    title: str,
    section: str,
    content: str,
    provenance: Mapping[str, object],
    inferred: bool = False,
) -> KnowledgeChunk:
    content_sha256 = hashlib.sha256(content.encode("utf-8")).hexdigest()
    identity = "\x00".join(
        (source_id, source_version, source_path, section, content_sha256)
    ).encode("utf-8")
    return KnowledgeChunk(
        chunk_id=hashlib.sha256(identity).hexdigest(),
        source_id=source_id,
        source_kind=source_kind,
        source_version=source_version,
        source_path=source_path,
        title=title,
        section=section,
        content=content,
        content_sha256=content_sha256,
        provenance=dict(provenance),
        inferred=inferred,
    )


def _decode_text(content: bytes) -> str:
    try:
        return content.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise DocumentError("document must be UTF-8 text") from exc


def _read_docx(content: bytes) -> str:
    try:
        from docx import Document
    except (ImportError, ModuleNotFoundError):
        raise DocumentDependencyUnavailable("python-docx is required for DOCX") from None
    try:
        document = Document(BytesIO(content))
    except Exception as exc:
        raise DocumentError(f"cannot parse DOCX: {exc}") from None
    blocks = [paragraph.text.strip() for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            blocks.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n\n".join(block for block in blocks if block)


def _read_pdf(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except (ImportError, ModuleNotFoundError):
        raise DocumentDependencyUnavailable("pypdf is required for PDF") from None
    try:
        reader = PdfReader(BytesIO(content))
        text = "\n\n".join((page.extract_text() or "").strip() for page in reader.pages)
    except Exception as exc:
        raise DocumentError(f"cannot parse PDF: {exc}") from None
    if not text.strip():
        raise UnsupportedDocument("unsupported_ocr")
    return text
